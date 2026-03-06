import re
import io
import warnings
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from lxml import etree
import os
import tempfile

# Wrapped formatting logic for backend usage

ORANGE = RGBColor(227, 108, 10)
BLUE   = RGBColor(18, 67, 149)
BLACK  = RGBColor(29, 29, 27)
WHITE  = RGBColor(255, 255, 255)

# No hardcoded heading dictionaries — detection is purely structure-based
# (numbering depth + bold formatting) so it works for ANY document.

# Explicit bullet marker characters: •, -, *, □, ▪ and Unicode variants
# Also catches plain 'o' if followed by space OR if directly followed by an uppercase letter/quote (e.g. "oAnd", `o"Shall"`)
BULLET_MARKER = re.compile(
    r'^(?:'
    r'[\u2022\u2023\u25CF\u25CB\u25AA\u25AB\u25A0\u25A1\u25B8\u25B9'
    r'\u2043\u2013\u2014\u2610\u25E6\u25C6\u25C7\uf0b7\uf0a7\uf076'
    r'\uf0d8\u00B7\-\*]\s*'
    r'|o(?=\s|[A-Z]|")\s*'
    r')'
)

# Numbered list prefix patterns → Word style mapping
# Order matters: check more specific patterns first
LIST_PREFIX_PATTERNS = [
    # Roman: (i), (ii), i), ii), i., ii.  etc.
    (re.compile(r'^\(([ivxlcdm]+)\)\s*', re.IGNORECASE), 'list_roman'),
    (re.compile(r'^([ivxlcdm]+)\)\s*', re.IGNORECASE),   'list_roman'),
    (re.compile(r'^([ivxlcdm]+)\.\s*', re.IGNORECASE),   'list_roman'),
    # Alphabetic: (a), a), a.  etc.
    (re.compile(r'^\(([a-zA-Z])\)\s*'),   'list_alpha'),
    (re.compile(r'^([a-zA-Z])\)\s*'),     'list_alpha'),
    (re.compile(r'^([a-zA-Z])\.\s+'),     'list_alpha'),   # require space to avoid matching e.g. "A4"
    # Numeric: 1., 1), (1)  etc.
    (re.compile(r'^\((\d+)\)\s*'),         'list_number'),
    (re.compile(r'^(\d+)\)\s*'),           'list_number'),
    (re.compile(r'^(\d+)\.\s*'),           'list_number'),
    # Bare numeric: "20List" or "20 List" (from flatten_numbering with no separator)
    (re.compile(r'^(\d+)\s+'),             'list_number'),   # number + space
    (re.compile(r'^(\d+)(?=[A-Z])'),       'list_number'),   # number directly followed by uppercase
]

# Map list types to Word styles
LIST_STYLE_MAP = {
    'list_number': 'List Number',
    'list_alpha':  'List Number 2',
    'list_roman':  'List Number 3',
}

CONTENT_WIDTH_CM = 21.0 - 2 * 2.54
CONTENT_WIDTH_EMU = int(CONTENT_WIDTH_CM * 360000)

MAX_WIDTH  = int(16 * 360000)
MAX_HEIGHT = int(20 * 360000)
MIN_SIZE   = int(1.0 * 360000)

MC_ALT = '{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent'
MC_CHOICE = '{http://schemas.openxmlformats.org/markup-compatibility/2006}Choice'
WPS_WSP = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}wsp'
WPG_WGP = '{http://schemas.microsoft.com/office/word/2010/wordprocessingGroup}wgp'
WPC_WPC = '{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}wpc'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

def set_font_xml(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def set_spacing(para, before=0, after=0, line_mult=1.05, auto_before=False, auto_after=False):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    for a in ['w:before','w:after','w:line','w:lineRule','w:beforeAutospacing','w:afterAutospacing']:
        k = qn(a)
        if k in sp.attrib:
            del sp.attrib[k]
    
    if auto_before:
        sp.set(qn('w:beforeAutospacing'), '1')
    else:
        sp.set(qn('w:before'), str(int(before * 20)))
        
    if auto_after:
        sp.set(qn('w:afterAutospacing'), '1')
    else:
        sp.set(qn('w:after'), str(int(after * 20)))
        
    if line_mult is not None:
        sp.set(qn('w:line'), str(int(line_mult * 240)))
        sp.set(qn('w:lineRule'), 'auto')

def set_shading(para, r, g, b):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')): pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    pPr.append(shd)

def add_run(para, text, font='Times New Roman', size=16, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color
    set_font_xml(run, font)
    return run

def is_shape_content(elem):
    if elem.findall('.//' + WPS_WSP): return True
    if elem.findall('.//' + WPG_WGP): return True
    if elem.findall('.//' + WPC_WPC): return True
    return False

def _set_para_mark_font(pPr, font_name, size_pt):
    """Set the paragraph mark font (rPr inside pPr) which controls list marker/number appearance."""
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    
    # Set font family
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    
    # Set font size (in half-points)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(size_pt * 2))  # Word uses half-points
    
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(size_pt * 2))

def apply_list_bullet(doc, para):
    """Apply real MS Word 'List Bullet' paragraph style.
    Never inserts bullet characters as text — uses proper Word list formatting."""
    try:
        para.style = doc.styles['List Bullet']
    except KeyError:
        # 'List Bullet' not in styles — create numbering XML manually
        pPr = para._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.insert(0, numPr)
    
    # Set indent for proper bullet alignment
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '720')    # 0.5 inch
    ind.set(qn('w:hanging'), '360') # hanging indent for bullet symbol
    
    # Override paragraph-level font to Times New Roman 16pt
    # This controls the bullet/number marker font
    _set_para_mark_font(pPr, 'Times New Roman', 16)

def strip_bullet_marker(full_text, runs):
    """Remove explicit bullet marker character from text and runs.
    Returns (cleaned_text, cleaned_runs, was_bullet)."""
    m = BULLET_MARKER.match(full_text)
    if not m:
        return full_text, runs, False
    
    prefix_len = len(m.group(0))
    cleaned_text = full_text[prefix_len:].strip()
    
    # Strip from runs too
    chars_left = prefix_len
    for rc in runs:
        if chars_left <= 0: break
        if len(rc['text']) <= chars_left:
            chars_left -= len(rc['text'])
            rc['text'] = ''
        else:
            rc['text'] = rc['text'][chars_left:].lstrip()
            chars_left = 0
    cleaned_runs = [rc for rc in runs if rc['text']]
    
    return cleaned_text, cleaned_runs, True

def strip_list_prefix(full_text, runs):
    """Detect and remove a numbered list prefix (numeric, alpha, roman).
    Returns (cleaned_text, cleaned_runs, list_type_or_None)."""
    for pattern, list_type in LIST_PREFIX_PATTERNS:
        m = pattern.match(full_text)
        if m:
            prefix_len = len(m.group(0))
            cleaned_text = full_text[prefix_len:].strip()
            
            # Don't match single letter followed by capital (likely a sentence, not a list)
            # e.g. "A computer is..." should NOT be treated as a list
            if list_type == 'list_alpha' and len(cleaned_text) > 0 and cleaned_text[0].isupper():
                # Only match if the alpha prefix is lowercase
                alpha_part = m.group(1) if m.lastindex else ''
                if alpha_part.isupper() and len(alpha_part) == 1:
                    continue  # Skip — likely just a sentence starting with "A" or "I"
            
            # Strip prefix from runs
            chars_left = prefix_len
            for rc in runs:
                if chars_left <= 0: break
                if len(rc['text']) <= chars_left:
                    chars_left -= len(rc['text'])
                    rc['text'] = ''
                else:
                    rc['text'] = rc['text'][chars_left:].lstrip()
                    chars_left = 0
            cleaned_runs = [rc for rc in runs if rc['text']]
            
            return cleaned_text, cleaned_runs, list_type
    
    return full_text, runs, None

def apply_list_style(doc, para, num_id):
    """Apply a real Word numbered list with a specific numId to a paragraph."""
    # We no longer rely on built-in styles; we apply formatting directly
    pPr = para._p.get_or_add_pPr()
    
    # Remove old numPr
    old_numPr = pPr.find(qn('w:numPr'))
    if old_numPr is not None:
        pPr.remove(old_numPr)
        
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_elem)
    pPr.insert(0, numPr)
    
    # Set indent for proper list alignment
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '720')    # 0.5 inch
    ind.set(qn('w:hanging'), '360') # hanging indent
    
    # Override paragraph-level font to Times New Roman 16pt
    _set_para_mark_font(pPr, 'Times New Roman', 16)

def setup_multilevel_heading_numbering(doc):
    """Create a multilevel list numbering definition that links Heading 1-4 to auto-numbering.
    Format: 1 → 1.1 → 1.1.1 → 1.1.1.1"""
    
    # Access or create the numbering part
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element
    
    # Define the abstract numbering with 4 levels
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), '10')  # Use ID 10 to avoid conflicts
    
    # Multi-level type
    multiLevelType = OxmlElement('w:multiLevelType')
    multiLevelType.set(qn('w:val'), 'multilevel')
    abstractNum.append(multiLevelType)
    
    # Level formats: 1, 1.1, 1.1.1, 1.1.1.1
    level_formats = [
        '',             # Level 0: hidden, so we manually print "CHAPTER-N"
        '%1.%2',        # Level 1: "1.1" 
        '%1.%2.%3',     # Level 2: "1.1.1"
        '%1.%2.%3.%4',  # Level 3: "1.1.1.1"
    ]
    
    heading_styles = ['Heading1', 'Heading2', 'Heading3', 'Heading4']
    # Per-level formatting: (size_pt, color_hex, bold)
    level_formatting = [
        (35, 'FFFFFF', True),    # H1: 35pt, white, bold
        (30, 'E36C0A', True),    # H2: 30pt, orange, bold
        (20, '124395', True),    # H3: 20pt, blue, bold
        (18, '1D1D1B', True),    # H4: 18pt, black, bold
    ]
    
    for i in range(4):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(i))
        
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'decimal')
        lvl.append(numFmt)
        
        # Link to heading style
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), heading_styles[i])
        lvl.append(pStyle)
        
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), level_formats[i])
        lvl.append(lvlText)
        
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        
        # Paragraph properties (Indentation & Tabs) to align wrapped text
        # If the space is smaller than the number width, Word pushes text to the NEXT default tab stop (creating a huge gap).
        # We start all numbers at 0" (left = hanging) and give them enough space to fit the text.
        if i == 0:    # H1: 35pt, e.g. "1" -> 0.5" space
            indent_left = "720"
            indent_hanging = "720"
        elif i == 1:  # H2: 30pt, e.g. "1.13" -> 0.75" space
            indent_left = "1080"
            indent_hanging = "1080"
        elif i == 2:  # H3: 20pt, e.g. "1.13.1" -> 0.75" space
            indent_left = "1080"
            indent_hanging = "1080"
        else:         # H4: 18pt, e.g. "1.13.1.1" -> 1.0" space
            indent_left = "1440"
            indent_hanging = "1440"
            
        pPr = OxmlElement('w:pPr')
        
        # Tab stop at the text indent position
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), indent_left)
        tabs.append(tab)
        pPr.append(tabs)
        
        # Indent
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), indent_left)
        ind.set(qn('w:hanging'), indent_hanging)
        pPr.append(ind)
        
        lvl.append(pPr)
        
        # Set number formatting to match heading text: font, size, color, bold
        size_pt, color_hex, is_bold = level_formatting[i]
        rPr = OxmlElement('w:rPr')
        
        # Font family
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:hint'), 'default')
        rPr.append(rFonts)
        
        # Font size (in half-points)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size_pt * 2))
        rPr.append(szCs)
        
        # Bold
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
        
        # Color
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), color_hex)
        rPr.append(color_elem)
        
        lvl.append(rPr)
        
        abstractNum.append(lvl)
    
        abstractNum.append(lvl)
    
    # Insert abstractNum before any existing w:num elements
    nums = numbering_elem.findall(qn('w:num'))
    if nums:
        nums[0].addprevious(abstractNum)
    else:
        numbering_elem.append(abstractNum)
    
    # Create the num element referencing abstractNumId=10
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '10')
    abstractNumId_elem = OxmlElement('w:abstractNumId')
    abstractNumId_elem.set(qn('w:val'), '10')
    num.append(abstractNumId_elem)
    numbering_elem.append(num)

def setup_body_lists(doc):
    """Create abstract numbering definitions for body lists: numeric, alpha, roman.
    Ids: 20=numeric, 21=alpha, 22=roman."""
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element
    
    formats = [
        ('20', 'decimal', '%1.', 'list_number'),   # 1., 2., 3.
        ('21', 'lowerLetter', '%1)', 'list_alpha'),# a), b), c)
        ('22', 'lowerRoman', '%1.', 'list_roman'), # i., ii., iii.
    ]
    
    for abs_id, num_fmt, lvl_text, list_type in formats:
        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), abs_id)
        
        multiLevelType = OxmlElement('w:multiLevelType')
        multiLevelType.set(qn('w:val'), 'hybridMultilevel')
        abstractNum.append(multiLevelType)
        
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), num_fmt)
        lvl.append(numFmt)
        
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), lvl_text)
        lvl.append(lvlText)
        
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '32') # 16pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '32')
        rPr.append(szCs)
        lvl.append(rPr)
        
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)
        
        abstractNum.append(lvl)
        
        nums = numbering_elem.findall(qn('w:num'))
        if nums:
            nums[0].addprevious(abstractNum)
        else:
            numbering_elem.append(abstractNum)

def _link_heading_to_numbering(para, heading_level):
    """Link a heading paragraph to the multilevel numbering list (numId=10)."""
    pPr = para._p.get_or_add_pPr()
    
    # Remove any existing numPr
    old_numPr = pPr.find(qn('w:numPr'))
    if old_numPr is not None:
        pPr.remove(old_numPr)
    
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(heading_level - 1))  # heading_level 1-based → ilvl 0-based
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '10')  # References the multilevel list we created
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.insert(0, numPr)

def heuristic_bullet_pass(items):
    """Post-extraction pass: detect consecutive short sentences that should be bullets.
    Triggers after headings containing 'you will be able to' or similar.
    Marks body items as is_bullet=True if they are short (<120 chars),
    start with a capital letter, and end with a period."""
    TRIGGER_PHRASES = [
        'you will be able to',
        'at the end of this lesson',
        'learning objectives',
        'after completing this',
    ]
    
    in_bullet_zone = False
    consecutive_count = 0
    candidate_start = -1
    
    for i, item in enumerate(items):
        ct = item['type']
        text = item.get('text', '').strip()
        
        # Check if a heading triggers the bullet zone
        if ct in ('h1', 'h2', 'h3', 'h4'):
            lower = text.lower()
            if any(phrase in lower for phrase in TRIGGER_PHRASES):
                in_bullet_zone = True
                consecutive_count = 0
                candidate_start = i + 1
            else:
                # A new heading that's not a trigger ends the zone
                if in_bullet_zone and consecutive_count >= 2:
                    # Mark the candidates as bullets
                    for j in range(candidate_start, candidate_start + consecutive_count):
                        if j < len(items) and items[j]['type'] == 'body':
                            items[j]['is_bullet'] = True
                in_bullet_zone = False
                consecutive_count = 0
            continue
        
        if in_bullet_zone and ct == 'body' and not item.get('is_bullet'):
            # Check heuristic: short, starts capital, ends period
            if (len(text) < 120 and len(text) > 3 
                    and text[0].isupper() and text.rstrip().endswith('.')):
                consecutive_count += 1
            else:
                # Break in pattern — mark what we have if >= 2 consecutive
                if consecutive_count >= 2:
                    for j in range(candidate_start, candidate_start + consecutive_count):
                        if j < len(items) and items[j]['type'] == 'body':
                            items[j]['is_bullet'] = True
                in_bullet_zone = False
                consecutive_count = 0
        elif ct != 'body':
            # Non-body item breaks the zone
            if in_bullet_zone and consecutive_count >= 2:
                for j in range(candidate_start, candidate_start + consecutive_count):
                    if j < len(items) and items[j]['type'] == 'body':
                        items[j]['is_bullet'] = True
            in_bullet_zone = False
            consecutive_count = 0
    
    # Handle end of document
    if in_bullet_zone and consecutive_count >= 2:
        for j in range(candidate_start, candidate_start + consecutive_count):
            if j < len(items) and items[j]['type'] == 'body':
                items[j]['is_bullet'] = True


def detect_chapter_headings(items):
    """
    Detect structural chapter boundaries by anchoring on the "Objectives" section.
    Rules:
    1. Find the paragraph containing "Objectives" (or "Learning Objectives").
    2. Look backwards to the immediately preceding text paragraph(s).
    3. Promote the preceding paragraph to an H1 chapter title with a page break.
    4. If there are TWO short paragraphs right before Objectives (e.g., "Chapter 1" then "Title"),
       combine them into a single H1 block.
    """
    # 1) Identify the index of the Objectives section(s)
    objective_indices = []
    for i, item in enumerate(items):
        text = item.get('text', '').strip().lower()
        if item['type'] in ('doc_title', 'h1', 'h2', 'h3', 'h4', 'h2_no_num', 'body'):
            if text in ('objectives', 'objectives:', 'learning objectives', 'learning objectives:'):
                objective_indices.append(i)

    # 2) For each Objectives section, look backwards to find the chapter title
    chapter_counter = 1
    for obj_idx in objective_indices:
        # Find the immediately preceding valid text paragraphs
        prev_texts = []
        for i in range(obj_idx - 1, -1, -1):
            if items[i].get('text', '').strip() and items[i]['type'] in ('doc_title', 'h1', 'h2', 'h3', 'h4', 'h2_no_num', 'body'):
                prev_texts.append(i)
                if len(prev_texts) == 2:
                    break
                    
        if not prev_texts:
            continue
            
        first_prev_idx = prev_texts[0]
        first_item = items[first_prev_idx]
        first_text = first_item.get('text', '').strip()
        first_word_count = len(first_text.split())

        # If there's a second preceding paragraph, check if it's a short chapter identifier (e.g. "CHAPTER 1" or "Unit 3")
        should_merge = False
        if len(prev_texts) == 2:
            second_prev_idx = prev_texts[1]
            second_item = items[second_prev_idx]
            second_text = second_item.get('text', '').strip()
            second_word_count = len(second_text.split())
            
            # If both lines are reasonably short (e.g. a number + a title), they should be merged
            if first_word_count <= 15 and second_word_count <= 10:
                should_merge = True

        if should_merge:
            # We merge the older paragraph (second_prev) and newer paragraph (first_prev)
            second_item = items[prev_texts[1]]
            first_item = items[prev_texts[0]]
            
            title_text = first_item['text'].strip()
            tag_text = second_item['text'].strip()
            
            # If tag text looks like "CHAPTER 1", replace it. Otherwise keep it.
            is_tag = False
            if any(char.isdigit() for char in tag_text) or any(w in tag_text.lower() for w in ('chapter', 'unit', 'lesson', 'module', 'part')):
                if len(tag_text.split()) <= 4:
                    is_tag = True
                    
            if is_tag:
                new_text = f"CHAPTER-{chapter_counter}\n{title_text}"
            else:
                new_text = f"CHAPTER-{chapter_counter}\n{tag_text}\n{title_text}"
            
            # The older paragraph becomes the H1 block
            second_item['type'] = 'h1'
            second_item['page_break'] = True
            second_item['text'] = new_text
            
            # The newer paragraph is flagged for deletion
            first_item['_merged'] = True
        else:
            # Just promote the single immediately preceding paragraph
            title_text = first_item['text'].strip()
            
            # Attempt to strip out leading identifiers like "Unit 3 - " or "Chapter 1: "
            title_text = re.sub(r'^(?:chapter|unit|lesson|module|part)?\s*\d+\s*[:\-\.]?\s*', '', title_text, flags=re.IGNORECASE)
            
            new_text = f"CHAPTER-{chapter_counter}\n{title_text}"
            first_item['type'] = 'h1'
            first_item['page_break'] = True
            first_item['text'] = new_text
            
        chapter_counter += 1

    # Filter out merged items
    items[:] = [it for it in items if not it.get('_merged')]

def extract_safe_image(src_doc, inline_elem):
    if is_shape_content(inline_elem): return None
    extent = inline_elem.find(qn('wp:extent'))
    if extent is None: return None
    cx, cy = int(extent.get('cx', '0')), int(extent.get('cy', '0'))
    if cx > MAX_WIDTH and cy > MAX_HEIGHT: return None
    if cx < MIN_SIZE and cy < MIN_SIZE: return None
    blips = inline_elem.findall('.//{%s}blip' % A_NS)
    if not blips: return None
    r_embed = blips[0].get('{%s}embed' % R_NS)
    if not r_embed: return None
    try:
        part = src_doc.part.related_parts[r_embed]
        data = part.blob
        if part.content_type and not part.content_type.startswith('image/'): return None
        return (data, cx, cy)
    except: return None



def set_table_borders(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

def flatten_numbering(input_path, output_path):
    """Uses MS Word via COM to permanently convert list numbers/bullets into plain text"""
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        try:
            # Must use absolute paths for COM
            abs_in = os.path.abspath(input_path)
            abs_out = os.path.abspath(output_path)
            doc = word.Documents.Open(abs_in)
            doc.ConvertNumbersToText()
            doc.SaveAs2(abs_out, 16) # wdFormatDocumentDefault
            doc.Close()
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
        return True
    except Exception as e:
        print(f"Failed to flatten numbering: {e}")
        return False

def format_document(input_path, output_path):
    print("PHASE 1: Extracting content from original raw file...")
    
    # Pre-process doc to flatten numbering and bullets to literal text
    flat_path = input_path + "_flat.docx"
    has_flat = flatten_numbering(input_path, flat_path)
    target_parse_path = flat_path if has_flat else input_path
    
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        raw = Document(target_parse_path)

    items = []
    first_text_found = False  # Track whether we've seen the first text paragraph (= title)

    for child in raw.element.body:
        tag = child.tag
        if tag == qn('w:p'):
            texts = [t.text for t in child.findall('.//' + qn('w:t')) if t.text]
            raw_full = ''.join(texts).replace('\u2028', ' ').replace('\t', ' ').replace('\n', ' ')
            full_text = re.sub(r' +', ' ', raw_full).strip()
            
            safe_images = []
            for d in child.findall('.//' + qn('w:drawing')):
                for inl in d.findall(qn('wp:inline')):
                    res = extract_safe_image(raw, inl)
                    if res: safe_images.append(res)
            for ac in child.findall('.//' + MC_ALT):
                choice = ac.find(MC_CHOICE)
                if choice is not None:
                    for d in choice.findall('.//' + qn('w:drawing')):
                        for inl in d.findall(qn('wp:inline')):
                            res = extract_safe_image(raw, inl)
                            if res: safe_images.append(res)
                
            if not full_text and not safe_images: continue
            
            # Skip stray page numbers (paragraphs that are purely digits)
            if full_text.strip().isdigit() and not safe_images:
                continue
                    
            runs = []
            for r in child.findall(qn('w:r')):
                rtext = ''.join(t.text for t in r.findall(qn('w:t')) if t.text)
                if not rtext: continue
                b, it = False, False
                rPr = r.find(qn('w:rPr'))
                if rPr is not None:
                    be = rPr.find(qn('w:b'))
                    if be is not None and be.get(qn('w:val')) in (None, 'true', '1', ''): b = True
                    ie = rPr.find(qn('w:i'))
                    if ie is not None and ie.get(qn('w:val')) in (None, 'true', '1', ''): it = True
                runs.append({'text': rtext, 'bold': b, 'italic': it})

            # --- Clean runs ---
            for rc in runs:
                rc['text'] = re.sub(r' +', ' ', rc['text'].replace('\u2028', ' ').replace('\t', ' ').replace('\n', ' '))
            for i in range(len(runs) - 1):
                if runs[i]['text'].endswith(' ') and runs[i+1]['text'].startswith(' '):
                    runs[i+1]['text'] = runs[i+1]['text'][1:]
            if runs:
                runs[0]['text'] = runs[0]['text'].lstrip()
                runs[-1]['text'] = runs[-1]['text'].rstrip()
            runs = [rc for rc in runs if rc['text']]
            # Rebuild full_text after cleaning runs
            if runs:
                full_text = ''.join(rc['text'] for rc in runs).strip()
            # ------------------

            ctype = 'body'
            lower_text = full_text.lower()
            
            # ── First text paragraph = Document Title ──
            if not first_text_found and full_text:
                ctype = 'doc_title'
                first_text_found = True
            else:
                # ── Hierarchical heading detection: count segments ──
                # 1.1.1.1 = h4, 1.1.1 = h3, 1.1 = h2, standalone chapter = h1
                heading_match = re.match(r'^(\d+(?:\.\d+)*)\s*(.*)', full_text)
                if heading_match:
                    num_part = heading_match.group(1)
                    segments = num_part.split('.')
                    seg_count = len(segments)
                    
                    if seg_count >= 2:  # X.Y, X.Y.Z, X.Y.Z.W → h2/h3/h4
                        if seg_count == 2:
                            ctype = 'h2'
                        elif seg_count == 3:
                            ctype = 'h3'
                        else:  # 4+
                            ctype = 'h4'
                        
                        # Strip the numeric prefix from text
                        stripped_title = heading_match.group(2).strip()
                        if stripped_title:
                            full_text = stripped_title
                            chars_to_strip = len(heading_match.group(0)) - len(heading_match.group(2))
                            for rc in runs:
                                if chars_to_strip <= 0: break
                                if len(rc['text']) <= chars_to_strip:
                                    chars_to_strip -= len(rc['text'])
                                    rc['text'] = ''
                                else:
                                    rc['text'] = rc['text'][chars_to_strip:].lstrip()
                                    chars_to_strip = 0
                            runs = [rc for rc in runs if rc['text']]
                    
                    elif seg_count == 1:  # X. Title — single-segment numbered heading
                        # Only treat as h2 if it looks like a section title:
                        #   - short (≤ 10 words)
                        #   - does NOT end with a sentence-ending period  
                        #   - title starts with an uppercase letter
                        _title_body = heading_match.group(2).strip()
                        _word_count = len(_title_body.split())
                        _ends_sentence = (_title_body.rstrip().endswith('.')
                                          and not _title_body.rstrip().endswith('...'))
                        _starts_upper = bool(_title_body) and _title_body[0].isupper()
                        if _title_body and _word_count <= 10 and not _ends_sentence and _starts_upper:
                            ctype = 'h2'
                            full_text = _title_body
                            # Strip the numeric prefix from runs too
                            chars_to_strip = len(heading_match.group(0)) - len(heading_match.group(2))
                            for rc in runs:
                                if chars_to_strip <= 0: break
                                if len(rc['text']) <= chars_to_strip:
                                    chars_to_strip -= len(rc['text'])
                                    rc['text'] = ''
                                else:
                                    rc['text'] = rc['text'][chars_to_strip:].lstrip()
                                    chars_to_strip = 0
                            runs = [rc for rc in runs if rc['text']]
                
                # Special keywords as H2 (no numbering)
                # Ensure 'objectives', 'summary', and practice questions are formatted consistently as standalone h2 headings.
                h2_keywords = ('introduction', 'introduction:',
                               'check your progress', 'check your progress:', 
                               'summary', 'summary:', 
                               'objectives', 'objectives:', 'ojectives', 'ojectives:',
                               'practise questions', 'practice questions',
                               'practise question', 'practice question',
                               'let us sum up', 'glossary')
                if ctype == 'body' and lower_text.strip() in h2_keywords:
                    ctype = 'h2_no_num'  # h2 formatting, but no heading numbering
                
                # Figure captions — only when text starts with "fig" keyword
                if ctype == 'body' and lower_text.startswith('fig'):
                    ctype = 'fig'
            
            # ── Explicit bullet marker detection ──
            is_bullet = False
            list_type = None  # 'list_number', 'list_alpha', 'list_roman', or None
            if ctype == 'body' and full_text:
                # Debug: print first char's Unicode codepoint
                first_char = full_text[0]
                safe_text = full_text[:50].encode('ascii', 'replace').decode('ascii')
                print(f"  BODY first_char=U+{ord(first_char):04X} text='{safe_text}'")
                
                # 1) Check for bullet markers first
                full_text, runs, is_bullet = strip_bullet_marker(full_text, runs)
                if is_bullet:
                    safe_stripped = full_text[:60].encode('ascii', 'replace').decode('ascii')
                    print(f"    -> BULLET STRIPPED: '{safe_stripped}'")
                
                # 2) If not a bullet, check for numbered list prefixes.
                #    Single-segment headings (N. Short Title) were already promoted
                #    to h2 above, so anything reaching here with N. is a real list item.
                if not is_bullet:
                    full_text, runs, list_type = strip_list_prefix(full_text, runs)
                    if list_type:
                        safe_stripped = full_text[:60].encode('ascii', 'replace').decode('ascii')
                        print(f"    -> LIST ({list_type}): '{safe_stripped}'")

            if safe_images and not full_text: ctype = 'img'

            items.append({'type': ctype, 'text': full_text, 'runs': runs, 'images': safe_images, 'is_bullet': is_bullet, 'list_type': list_type})
            
        elif tag == qn('w:tbl'):
            rows = []
            for tr in child.findall(qn('w:tr')):
                cells = []
                for tc in tr.findall(qn('w:tc')):
                    ct = ''.join(t.text for t in tc.findall('.//' + qn('w:t')) if t.text).strip()
                    cells.append(ct)
                rows.append(cells)
            if rows: items.append({'type': 'table', 'rows': rows})

    # ── Structural chapter detection pass ──
    detect_chapter_headings(items)

    # ── Heuristic bullet pass: detect short sentence lists ──
    heuristic_bullet_pass(items)
    bullet_count = sum(1 for it in items if it.get('is_bullet'))
    print(f"  Total bullets detected: {bullet_count}")

    print("PHASE 2: Creating new document...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    # Moderate margins: Top/Bottom 2.54cm, Left/Right 1.91cm
    sec.top_margin, sec.bottom_margin = Cm(2.54), Cm(2.54)
    sec.left_margin, sec.right_margin = Cm(1.91), Cm(1.91)

    for hf in [sec.header, sec.footer]:
        for child in list(hf._element): hf._element.remove(child)
        etree.SubElement(hf._element, qn('w:p'))

    # Set up multilevel heading numbering (links Heading 1-4 to auto-numbering)
    setup_multilevel_heading_numbering(doc)
    setup_body_lists(doc)

    print("PHASE 3: Rebuilding...")
    
    # Track list state to handle list restarting
    active_lists = {}
    next_num_id_counter = 11  # 10 is used by headings
    
    for item in items:
        ct = item['type']
        
        # ──────────────────────────────────────────────────────────────────────
        # NOTE: We do NOT clear active_lists on headings any more.
        # Clearing caused every list that followed a heading to restart at "1."
        # Instead we only create a new numId the first time a list_type is seen.
        # A list restarts naturally only when normal body paragraphs interrupt it
        # (handled below at the else: active_lists.clear() branch).
        # ──────────────────────────────────────────────────────────────────────
        
        if ct in ['h1', 'doc_title']:
            p = doc.add_paragraph()
            
            # Only link actual chapters to the multilevel numbering (so numbering resets properly)
            if ct == 'h1':
                try:
                    p.style = doc.styles['Heading 1']
                except KeyError:
                    pass
                
            if item.get('page_break'):
                p.paragraph_format.page_break_before = True
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_shading(p, 227, 108, 10)
            
            parts = item['text'].upper().split('\n')
            for i, part_text in enumerate(parts):
                if i > 0:
                    p.add_run().add_break()
                add_run(p, part_text, 'Montserrat', 35, bold=True, color=WHITE)
            set_spacing(p, auto_before=True, auto_after=True, line_mult=1.05)
            
            # Link to multilevel numbering so it resets H2/H3/H4 prefixes (e.g. 1.1 -> 2.1)
            if ct == 'h1':
                _link_heading_to_numbering(p, 1)
            
        elif ct in ['h2', 'h3', 'h4', 'h2_no_num']:
            # Apply real Word Heading style for auto-numbering
            actual_level = 'h2' if ct == 'h2_no_num' else ct
            heading_level = {'h2': 2, 'h3': 3, 'h4': 4}[actual_level]
            style_name = f'Heading {heading_level}'
            
            p = doc.add_paragraph()
            try:
                p.style = doc.styles[style_name]
            except KeyError:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            color = ORANGE if actual_level=='h2' else (BLUE if actual_level=='h3' else BLACK)
            size = 30 if actual_level=='h2' else (20 if actual_level=='h3' else 18)
            
            fmt_text = item['text'].title()
            
            add_run(p, fmt_text, 'Times New Roman', size, bold=True, color=color)
            set_spacing(p, auto_before=True, auto_after=True, line_mult=1.05)
            
            # Link to multilevel numbering ONLY for numbered headings
            if ct != 'h2_no_num':
                _link_heading_to_numbering(p, heading_level)
            
        elif ct == 'fig':
            active_lists.clear()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, item['text'], 'Times New Roman', 14, bold=True, italic=True, color=BLACK)
            set_spacing(p, before=0, after=10, line_mult=1.05)
            
        elif ct == 'img':
            active_lists.clear()
            for data, w, h in item.get('images', []):
                try:
                    if w > 0 and w != CONTENT_WIDTH_EMU:
                        ratio = CONTENT_WIDTH_EMU / w
                        w, h = CONTENT_WIDTH_EMU, int(h * ratio)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(io.BytesIO(data), width=Emu(w))
                    set_spacing(p, before=0, after=10, line_mult=1.05)
                except: pass
                
        elif ct == 'body':
            for data, w, h in item.get('images', []):
                try:
                    if w > 0 and w != CONTENT_WIDTH_EMU:
                        ratio = CONTENT_WIDTH_EMU / w
                        w, h = CONTENT_WIDTH_EMU, int(h * ratio)
                    pi = doc.add_paragraph()
                    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pi.add_run().add_picture(io.BytesIO(data), width=Emu(w))
                    set_spacing(pi, before=0, after=10, line_mult=1.05)
                except: pass
            if item['text']:
                p = doc.add_paragraph()
                
                # Apply real Word list formatting if detected
                is_bullet = item.get('is_bullet')
                list_type = item.get('list_type')
                
                if is_bullet:
                    apply_list_bullet(doc, p)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif list_type:
                    # If new list sequence for this type, create a new num XML linked to the abstract list
                    if list_type not in active_lists:
                        current_num_id = next_num_id_counter
                        next_num_id_counter += 1
                        active_lists[list_type] = current_num_id
                        
                        abs_id = {'list_number': '20', 'list_alpha': '21', 'list_roman': '22'}[list_type]
                        numbering_elem = doc.part.numbering_part._element
                        num = OxmlElement('w:num')
                        num.set(qn('w:numId'), str(current_num_id))
                        
                        abstractNumId_elem = OxmlElement('w:abstractNumId')
                        abstractNumId_elem.set(qn('w:val'), abs_id)
                        num.append(abstractNumId_elem)
                        
                        # Add lvlOverride to FORCE Word to restart at 1
                        lvlOverride = OxmlElement('w:lvlOverride')
                        lvlOverride.set(qn('w:ilvl'), '0')
                        startOverride = OxmlElement('w:startOverride')
                        startOverride.set(qn('w:val'), '1')
                        lvlOverride.append(startOverride)
                        num.append(lvlOverride)
                        
                        numbering_elem.append(num)
                        
                    apply_list_style(doc, p, active_lists[list_type])
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # Normal paragraph breaks the list sequence
                    active_lists.clear()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                if item['runs']:
                    for r in item['runs']:
                        add_run(p, r['text'], 'Times New Roman', 16, bold=r['bold'], italic=r['italic'], color=BLACK)
                else:
                    add_run(p, item['text'], 'Times New Roman', 16, color=BLACK)
                set_spacing(p, before=0, after=6, line_mult=1.05)
                
        elif ct == 'table':
            active_lists.clear()
            rows = item['rows']
            max_cols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=max_cols)
            try:
                tbl.style = 'Table Grid'
            except: pass
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < max_cols:
                        cell = tbl.cell(ri, ci)
                        cell.paragraphs[0].text = ''
                        p = cell.paragraphs[0]
                        add_run(p, cell_text, 'Times New Roman', 12, bold=(ri==0), color=BLACK)
                        set_spacing(p, before=0, after=0, line_mult=1.0)

    for t in doc.tables:
        set_table_borders(t)

    print(f"Saving to {output_path}")
    doc.save(output_path)
    
    # Cleanup temp flattened file
    if has_flat:
        try: os.remove(flat_path)
        except: pass
        
    print("DONE!")

