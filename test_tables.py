from docx import Document
import os

for f in os.listdir('uploads'):
    if not f.endswith('flat.docx') and not f.endswith('6.7_Autocad.docx'): 
        continue
    
    path = os.path.join('uploads', f)
    try:
        doc = Document(path)
        print(f"\n--- {f} ---")
        
        # Check tables
        for i, tbl in enumerate(doc.tables):
            for row in tbl.rows:
                for cell in row.cells:
                    if 'UNIT' in cell.text.upper() and ('1' in cell.text or '2' in cell.text):
                        print(f"Table {i} Cell: {repr(cell.text)}")
                        
        # Check paragraphs
        for i, p in enumerate(doc.paragraphs):
            if 'UNIT' in p.text.upper() and ('1' in p.text or '2' in p.text):
                print(f"Para {i}: {repr(p.text)}")
                
    except Exception as e:
        print(f"Failed {f}: {e}")
