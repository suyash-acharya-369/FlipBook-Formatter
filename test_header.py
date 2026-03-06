from docx import Document
from docx.oxml.ns import qn
import sys

# Load the template
path = r"d:\DESKTOP\FORMATTING\Formatted Sample word file- DCA FUNDAMENTALS OF COMPUTER AND INFORMATION TECHNOLOGY - updated (03) (2).docx"
doc = Document(path)

print("--- HEADER INSPECTION ---")
for i, section in enumerate(doc.sections):
    print(f"Section {i} Header:")
    for j, p in enumerate(section.header.paragraphs):
        print(f"  Para {j}: '{p.text}'")
        for k, run in enumerate(p.runs):
            print(f"    Run {k}: '{run.text}' - bold:{run.bold}, italic:{run.italic}")

    if section.first_page_header:
        print(f"Section {i} First Page Header:")
        for j, p in enumerate(section.first_page_header.paragraphs):
            print(f"  Para {j}: '{p.text}'")
    if section.even_page_header:
        print(f"Section {i} Even Page Header:")
        for j, p in enumerate(section.even_page_header.paragraphs):
            print(f"  Para {j}: '{p.text}'")

        
print("--- SHAPES/TEXTBOXES IN HEADER ---")
for i, section in enumerate(doc.sections):
    for p in section.header.paragraphs:
        for run in p.runs:
            for tbox in run._element.xpath('.//w:t'):
                print(f"    Found text inside shape/element: {tbox.text}")

