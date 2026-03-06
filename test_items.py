import sys
import os
from docx import Document

print("--- SCANNING FOR ANY VARIATION OF UNIT TITLES ---")
for f in os.listdir("uploads"):
    if not f.endswith('.docx') or f.endswith('flat.docx') or not 'Automobile' in f:
        continue
    
    path = os.path.join("uploads", f)
    try:
        doc = Document(path)
        
        print("\n--- FIRST 50 PARAGRAPHS ---")
        for i, p in enumerate(doc.paragraphs[:50]):
            text = p.text.strip()
            if text:
                print(f"[{i}] {repr(text)}")
                
        print("\n--- FIRST 5 TABLES ---")
        for i, tbl in enumerate(doc.tables[:5]):
            print(f"TABLE {i}")
            for r, row in enumerate(tbl.rows):
                for c, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        print(f"  [{r},{c}] {repr(text)}")
                        
    except Exception as e:
        print(f"Could not read {f}: {e}")
