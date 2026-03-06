from docx import Document
from docx.oxml.ns import qn
import sys

# Load the template
path = r"d:\DESKTOP\FORMATTING\Formatted Sample word file- DCA FUNDAMENTALS OF COMPUTER AND INFORMATION TECHNOLOGY - updated (03) (2).docx"
doc = Document(path)

# Print initial info
print("Initial headers:")
for i, section in enumerate(doc.sections):
    print(f"Sec {i} header p count:", len(section.header.paragraphs))
print("Initial paragraphs:", len(doc.paragraphs))

# Strip the body
for element in list(doc.element.body):
    if element.tag != qn('w:sectPr'):
        doc.element.body.remove(element)

print("After strip paragraphs:", len(doc.paragraphs))

# Add a test paragraph
doc.add_paragraph("This is the new body.")

# Save to test output
out_path = "test_template_clear.docx"
doc.save(out_path)
print("Saved to", out_path)
